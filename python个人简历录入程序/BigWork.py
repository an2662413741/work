import tkinter.ttk
from tkinter import *
from docx import Document


# 个人简历模板1
def createword1(name, sex, zc, whcd, xxjl, jcxx):
    name = name
    sex = sex
    zc = zc
    whcd = whcd
    xxjl = xxjl
    jcxx = jcxx
    print(name, sex, zc, whcd, xxjl, jcxx)
    # 生成word文档
    document = Document()

    # 添加文章大标题，普通段落
    document.add_heading('个人简介', level=0)
    document.add_paragraph('姓         名：' + name, style='List Bullet')
    document.add_paragraph('性         别：' + sex, style='List Bullet')
    document.add_paragraph('专         业：' + zc, style='List Bullet')
    document.add_paragraph('文化程度：' + whcd, style='List Bullet')

    # 添加一级标题
    document.add_heading('学习经历：', level=1)
    document.add_paragraph(xxjl, style=None)

    # 添加一级列表
    document.add_heading('奖惩信息', level=1)
    document.add_paragraph(jcxx, style=None)

    # 保存文档
    document.save(name+'个人简历1.docx')


# 个人简历模板2
def createword2(name, sex, zc, whcd, xxjl, jcxx):
    name = name
    sex = sex
    zc = zc
    whcd = whcd
    xxjl = xxjl
    jcxx = jcxx
    print(name, sex, zc, whcd, xxjl, jcxx)
    # 生成word文档
    document = Document()

    # 添加文章大标题，普通段落
    document.add_heading('个人简介', level=0)
    table = document.add_table(rows=12, cols=4)

    table.cell(2, 0).merge(table.cell(2, 3))
    table.cell(3, 0).merge(table.cell(6, 3))
    table.cell(7, 0).merge(table.cell(7, 3))
    table.cell(8, 0).merge(table.cell(11, 3))

    table.cell(0, 0).text = "姓名"
    table.cell(0, 1).text = name
    table.cell(0, 2).text = "性别"
    table.cell(0, 3).text = sex

    table.cell(1, 0).text = "专业"
    table.cell(1, 1).text = zc
    table.cell(1, 2).text = "文化程度"
    table.cell(1, 3).text = whcd

    table.cell(2, 0).text = '学习经历'
    table.cell(3, 0).text = xxjl

    table.cell(7, 0).text = '奖惩信息'
    table.cell(8, 0).text = jcxx

    # 保存文档
    document.save(name+'个人简历2.docx')


# 获取值并创建个人简历.docx
def getxx():
    e1 = entry1.get()  # 获取Entry的内容
    s1 = sex.get()
    v1 = var1.get()
    v2 = var2.get()
    t1 = xxjl.get('0.0','end')
    t2 = jcxx.get('0.0','end')
    j1 = jlmb.get()
    print(e1, s1, v1, v2, t1, t2, j1)
    if j1 == 1:
        createword1(e1, s1, v1, v2, t1, t2)
    elif j1 == 2:
        createword2(e1, s1, v1, v2, t1, t2)


win = Tk()
win.geometry("1000x500")
win.title('个人简历录入系统')


fm1 = Frame(win)
lable11 = Label(fm1, text="姓    名：", font=('宋体', 24))
lable11.pack(side=LEFT)
entry1 = tkinter.Entry(fm1, bd=10)
entry1.pack(side=LEFT)
fm1.pack(side=TOP)


fm2 = Frame(win)
lable12 = Label(fm2, text="性    别：", font=('宋体', 24)).pack(side=LEFT)
sex = StringVar()
sex1 = Radiobutton(fm2, text="男", value='男', variable=sex, font=('宋体', 24)).pack()
sex2 = Radiobutton(fm2, text="女", value='女', variable=sex, font=('宋体', 24)).pack()
fm2.pack(side=TOP)


fm3 = Frame(win)
lable13 = Label(fm3, text="专    业：", font=('宋体', 24)).pack(side=LEFT)
var1 = StringVar()
zy = tkinter.ttk.Combobox(fm3, textvariable=var1, value=('计算机', '电子', '工商', '人事')).pack(padx=5, pady=10)
fm3.pack(side=TOP)


fm4 = Frame(win)
lable14 = Label(fm4, text="文化程度：", font=('宋体', 24)).pack(side=LEFT)
var2 = StringVar()
whcd = tkinter.ttk.Combobox(fm4, textvariable=var2, value=('大专', '本科', '硕士', '博士')).pack(padx=5, pady=10)
fm4.pack(side=TOP)


fm5 = Frame(win)
lable15 = Label(fm5, text="学习经历：", font=('宋体', 24)).pack(side=LEFT)
xxjl = Text(fm5, width='40', height='3')
xxjl.pack()
fm5.pack(side=TOP)


fm6 = Frame(win)
lable16 = Label(fm6, text="奖惩信息：", font=('宋体', 24)).pack(side=LEFT)
jcxx = Text(fm6, width='40', height='3')
jcxx.pack()
fm6.pack(side=TOP)


fm7 = Frame(win)
lable14 = Label(fm7, text="简历模板：", font=('宋体', 24)).pack(side=LEFT)
jlmb = IntVar()
jlmb1 = Radiobutton(fm7, text="模板一", value=1, variable=jlmb, font=('宋体', 24)).pack()
jlmb2 = Radiobutton(fm7, text="模板二", value=2, variable=jlmb, font=('宋体', 24)).pack()
fm7.pack(side=TOP)


fm8 = Frame(win)
button = Button(fm8, text='打印', font=('宋体', 24), command=getxx).pack()
fm8.pack(side=TOP)


win.mainloop()